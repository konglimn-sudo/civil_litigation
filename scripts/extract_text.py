#!/usr/bin/env python3
"""可选的本机文字抽取加速器（被 server.mjs 的 extractWithPython 调用）。

依赖 PyMuPDF(fitz) 与 python-docx,效果优于零依赖的 ocr.mjs(尤其中文 PDF/DOCX)。
约定:命令行第一个参数为文件路径;成功时向 stdout 打印
  {"text", "method", "pages", "warnings"} JSON,失败时打印 {"error"} 并以非 0 退出。
本机处理,数据不出本地。未安装本脚本依赖时,server.mjs 会自动回退到 ocr.mjs。
"""
import json
import os
import subprocess
import sys
import tempfile
from pathlib import Path


def run_tesseract(image_path):
    """对单张图片调用系统 tesseract,返回识别文本。"""
    command = [
        os.environ.get("TESSERACT_BIN", "tesseract"),  # 可执行文件(允许环境变量覆盖)
        str(image_path),
        "stdout",                                       # 结果输出到标准输出
        "-l",
        "chi_sim+eng",                                  # 简体中文 + 英文
        "--psm",
        "6",                                            # 假定为统一文本块
    ]
    result = subprocess.run(command, capture_output=True, text=True, timeout=120)
    if result.returncode != 0:                          # 非 0 退出视为失败
        raise RuntimeError(result.stderr.strip() or "Tesseract OCR failed")
    return result.stdout.strip()


def extract_pdf(file_path):
    """PDF:先取文本层;文本过少则按扫描件逐页栅格化 OCR。"""
    import fitz  # PyMuPDF,延迟导入以便缺失时只影响 PDF 分支

    document = fitz.open(file_path)
    text_parts = [page.get_text("text").strip() for page in document]  # 逐页取文本层
    text = "\n\n".join(part for part in text_parts if part)
    if len(text.strip()) >= 80:                         # 文本层足够长=数字版 PDF
        return text, "pdf-text", len(document), []

    # 否则视为扫描件:逐页渲染成图片再 OCR。
    warnings = []
    ocr_parts = []
    page_limit = min(len(document), 30)                 # 最多处理前 30 页
    with tempfile.TemporaryDirectory(prefix="hengfa-ocr-") as temp_dir:  # 临时目录,退出即删
        for index in range(page_limit):
            page = document[index]
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)  # 2 倍放大提升 OCR 准确率
            image_path = Path(temp_dir) / f"page-{index + 1}.png"
            pixmap.save(image_path)
            page_text = run_tesseract(image_path)       # 对该页 OCR
            if page_text:
                ocr_parts.append(f"[第 {index + 1} 页]\n{page_text}")
    if len(document) > page_limit:                      # 超出页数上限时给出提示
        warnings.append(f"扫描件超过 {page_limit} 页，仅处理前 {page_limit} 页")
    return "\n\n".join(ocr_parts), "pdf-ocr", len(document), warnings


def extract_docx(file_path):
    """DOCX:用 python-docx 取段落与表格文本。"""
    from docx import Document  # 延迟导入

    document = Document(file_path)
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]  # 段落
    for table in document.tables:                       # 表格:逐行用制表符连接单元格
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                parts.append("\t".join(values))
    return "\n".join(parts), "docx-text", None, []


def extract_plain(file_path):
    """纯文本:依次尝试 utf-8 / gb18030 / utf-16 解码。"""
    raw = Path(file_path).read_bytes()
    for encoding in ("utf-8", "gb18030", "utf-16"):
        try:
            return raw.decode(encoding), f"plain-{encoding}", None, []
        except UnicodeDecodeError:
            continue                                    # 该编码失败则换下一种
    # 都失败时用替换符兜底,并标注警告。
    return raw.decode("utf-8", errors="replace"), "plain-recovered", None, ["文本编码无法完整识别"]


def main():
    """按扩展名分派到对应抽取函数,并把结果以 JSON 打印到 stdout。"""
    file_path = Path(sys.argv[1]).resolve()             # 第一个参数=文件路径
    suffix = file_path.suffix.lower()                   # 扩展名(小写)
    if suffix == ".pdf":
        text, method, pages, warnings = extract_pdf(file_path)
    elif suffix == ".docx":
        text, method, pages, warnings = extract_docx(file_path)
    elif suffix in {".txt", ".md", ".csv", ".json"}:
        text, method, pages, warnings = extract_plain(file_path)
    elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}:
        text, method, pages, warnings = run_tesseract(file_path), "image-ocr", 1, []
    else:
        raise RuntimeError(f"Unsupported file type: {suffix}")  # 未知类型由 except 收敛为 JSON 错误
    print(json.dumps({"text": text.strip(), "method": method, "pages": pages, "warnings": warnings}, ensure_ascii=False))


if __name__ == "__main__":
    try:
        main()
    except Exception as error:                          # 任何异常都转成 {"error"} 供 Node 端解析
        print(json.dumps({"error": str(error)}, ensure_ascii=False))
        sys.exit(1)
